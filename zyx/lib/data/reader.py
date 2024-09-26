from concurrent.futures import ThreadPoolExecutor
import multiprocessing as mp
from loguru import logger
from pathlib import Path
import mimetypes
import json
from typing import Dict, List, Optional, Literal, Union, Type
from contextlib import suppress
import xml.etree.ElementTree as ET
import PyPDF2
import docx
import openpyxl
import csv
import requests
from urllib.parse import urlparse


from ..types.document import Document


OutputType = Literal["markdown", "text", "json"]
OutputFormat = Literal["document"]


import requests
from urllib.parse import urlparse

def read(
    path: Union[str, Path],
    output : Union[Type[str], OutputFormat] = "document",
    target: OutputType = "text",
    verbose: bool = False
) -> Union[Document, List[Document], str]:
    """
    Reads either a file or a directory and returns the content.
    """
    path = _download_if_url(path)
    path = Path(path)
    if path.is_file():
        return _read_single_file(
            path = path,
            output = output,
            target = target,
            verbose = verbose
        )
    elif path.is_dir():
        with ThreadPoolExecutor(max_workers=mp.cpu_count()) as executor:
            futures = [executor.submit(_read_single_file, file, output, target, verbose)
                       for file in path.glob('*') if file.is_file()]
            results = [future.result() for future in futures]
        return [result for result in results if result is not None]
    else:
        raise ValueError(f"Invalid path: {path}")

def _download_if_url(path: Union[str, Path]) -> Union[str, Path]:
    """
    Downloads the file if the path is a URL and returns the local file path.
    """
    if isinstance(path, str) and urlparse(path).scheme in ('http', 'https'):
        response = requests.get(path)
        response.raise_for_status()

        # Extract filename from URL or use a default one
        filename = Path(urlparse(path).path).name
        if not filename:
            filename = 'downloaded_file'

        # Try to get extension from Content-Type header if not in filename
        if not Path(filename).suffix:
            content_type = response.headers.get('Content-Type')
            if content_type:
                extension = mimetypes.guess_extension(content_type.split(';')[0])
                if extension:
                    filename += extension

        local_path = Path("/tmp") / filename
        with open(local_path, 'wb') as file:
            file.write(response.content)
        return local_path
    return path


def _read_single_file(
    path: Union[str, Path],
    target: OutputType = "text",
    output : Union[Type[str], OutputFormat] = "document",
    verbose: bool = False
) -> Optional[Document]:
    """
    Reads a single file and returns its content based on the target format.
    """
    path = Path(path)
    mime_type, _ = mimetypes.guess_type(str(path))

    if mime_type is None:
        # Attempt to detect file type by reading the first few bytes
        with open(path, 'rb') as f:
            header = f.read(5)
            if header == b'%PDF-':
                mime_type = 'application/pdf'
            elif header[:2] == b'PK':
                # Possible DOCX or XLSX (which are zip files)
                # Read the content types in the file to distinguish
                f.seek(0)
                file_content = f.read()
                if b'word/' in file_content:
                    mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif b'xl/' in file_content:
                    mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            elif header.startswith(b'\xEF\xBB\xBF') or header.startswith(b'\xFE\xFF') or header.startswith(b'\xFF\xFE'):
                # Possible text file with BOM
                mime_type = 'text/plain'
            else:
                # Default to binary if no match
                mime_type = 'application/octet-stream'

    try:
        content = None
        match mime_type:
            case 'application/pdf':
                content = _read_pdf(path)
            case 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                content = _read_docx(path)
            case 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                content = _read_xlsx(path)
            case 'text/csv':
                content = _read_csv(path)
            case _ if mime_type and mime_type.startswith('text/'):
                content = _read_text(path)
            case 'application/xml':
                content = _read_xml(path)
            case _:
                content = _read_binary(path)

        if verbose:
            logger.info(f"Read {path.name} as {mime_type or 'BINARY'}")

        if target == "markdown":
            content = _to_markdown(content)
        elif target == "json":
            content = json.dumps(content)

        metadata = {
            "file_name": path.name,
            "file_type": mime_type or "unknown",
            "file_size": path.stat().st_size
        }

        if output == "document":
            return Document(content=content, metadata=metadata)
        elif output == str:
            return content

        return Document(content=content, metadata=metadata)

    except Exception as e:
        if verbose:
            logger.error(f"Error reading file {path}: {str(e)}")
        return None


def _read_pdf(path: Path) -> str:
    """
    Extracts text from a PDF, including proper formatting for tables and paragraphs.
    """
    try:
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                extracted_text = page.extract_text() or ""
                # Additional handling for simple table formatting
                # We'll assume columns are separated by enough space to infer table-like structures
                text += _format_pdf_text(extracted_text)
        return text.strip()
    except Exception as e:
        logger.error(f"Error reading PDF {path}: {str(e)}")
        return ""

def _format_pdf_text(extracted_text: str) -> str:
    """
    Post-process PDF extracted text to ensure proper formatting for tables, paragraphs, etc.
    """
    # Insert custom logic to handle common cases like columns and tables, for example:
    # - Replace multiple spaces with a ' | ' for column-separated data in tables
    # - Handle paragraph breaks or other structures

    # This is a placeholder and can be adjusted based on the specifics of the PDF content
    lines = extracted_text.split("\n")
    formatted_lines = []
    for line in lines:
        if "  " in line:  # Simple heuristic for columns in tables
            formatted_lines.append(" | ".join(line.split()))
        else:
            formatted_lines.append(line)
    return "\n".join(formatted_lines)

def _read_docx(path: Path) -> str:
    """
    Reads text from a DOCX file, including tables.
    """
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Handle tables in the DOCX file
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_data))  # Simple table formatting with '|'

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading DOCX {path}: {str(e)}")
        return ""

def _read_xlsx(path: Path) -> List[List[str]]:
    """
    Reads an XLSX file and returns its data, including table formatting.
    """
    try:
        with openpyxl.load_workbook(path, read_only=True) as workbook:
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows():
                row_data = [str(cell.value) if cell.value is not None else "" for cell in row]
                data.append(row_data)
            return data
    except Exception as e:
        logger.error(f"Error reading XLSX {path}: {str(e)}")
        return []

def _read_csv(path: Path) -> List[List[str]]:
    """
    Reads CSV data.
    """
    try:
        with open(path, 'r', newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            return list(reader)
    except Exception as e:
        logger.error(f"Error reading CSV {path}: {str(e)}")
        return []

def _read_text(path: Path) -> str:
    """
    Reads plain text files.
    """
    try:
        with open(path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading text file {path}: {str(e)}")
        return ""

def _read_xml(path: Path) -> Dict:
    """
    Reads XML files and converts them to a dictionary.
    """
    try:
        tree = ET.parse(path)
        root = tree.getroot()
        return _element_to_dict(root)
    except Exception as e:
        logger.error(f"Error reading XML {path}: {str(e)}")
        return {}

def _read_binary(path: Path) -> str:
    return f"Binary file: {path.name}"

def _element_to_dict(element):
    """
    Recursively converts XML elements to a dictionary.
    """
    result = {}
    for child in element:
        if len(child) == 0:
            result[child.tag] = child.text
        else:
            result[child.tag] = _element_to_dict(child)
    return result

def _to_markdown(content: Union[str, List, Dict]) -> str:
    """
    Converts the content into markdown format.
    """
    if isinstance(content, str):
        return content
    elif isinstance(content, list):
        return "\n".join([" | ".join(row) for row in content])
    elif isinstance(content, dict):
        return "\n".join([f"**{k}**: {v}" for k, v in content.items()])
    else:
        return str(content)
