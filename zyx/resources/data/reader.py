try:
    from concurrent.futures import ThreadPoolExecutor
    import multiprocessing as mp
    from pathlib import Path
    import mimetypes
    import json
    from typing import Dict, List, Optional, Literal, Union, Type
    import xml.etree.ElementTree as ET
    import PyPDF2
    import docx
    import openpyxl
    import csv
    import requests
    from urllib.parse import urlparse
    import uuid
    import os

    from ...lib.types.document import Document
    from ...lib.utils.logger import get_logger
except ImportError:
    import os

    print(
        "The [bold]`zyx(data)`[/bold] data extension is required to use this module. Install it?"
    )
    if input("Install? (y/n)") == "y":
        os.system("pip install 'zyx[data]'")
    else:
        print("Exiting...")
        exit(1)

logger = get_logger("reader")

OutputType = Literal["markdown", "text", "json"]
OutputFormat = Literal["document", "json"]

import requests
from urllib.parse import urlparse


def read(
    path: Union[str, Path, List[Union[str, Path]]],
    output: Union[Type[str], OutputFormat] = "document",
    target: OutputType = "text",
    verbose: bool = False,
    workers: Optional[int] = None,
) -> Union[Document, List[Document], str, Dict, List[Dict]]:
    """
    Reads either a file, a directory, or a list of files and returns the content.

    Example:
        ```python
        read("path/to/file.pdf")
        # Document(content="...", metadata={"file_name": "file.pdf", "file_type": "application/pdf", "file_size": 123456})
        ```

    Args:
        path: Union[str, Path, List[Union[str, Path]]]: The path to read.
        output: Union[Type[str], OutputFormat]: The output format.
        target: OutputType: The output type.
        verbose: bool: Whether to print verbose output.
        workers: Optional[int]: The number of workers to use for reading.

    Returns:
        Union[Document, List[Document], str, Dict, List[Dict]]: The content.
    """
    if isinstance(path, list):
        paths = [_download_if_url(p) for p in path]
    else:
        paths = [_download_if_url(path)]

    paths = [Path(p) for p in paths]

    try:
        if len(paths) == 1 and paths[0].is_file():
            result = _read_single_file(
                path=paths[0], output=output, target=target, verbose=verbose
            )
            if output == "json":
                return result  # Directly return the result if it's JSON
            return result
        else:
            with ThreadPoolExecutor(max_workers=workers or mp.cpu_count()) as executor:
                futures = [
                    executor.submit(_read_single_file, file, output, target, verbose)
                    for p in paths
                    for file in (p.glob("*") if p.is_dir() else [p])
                    if file.is_file()
                ]
                results = [future.result() for future in futures]
            if output == "json":
                return results  # Directly return the list of results if it's JSON
            return [result for result in results if result is not None]
    finally:
        # Cleanup temporary files
        for p in paths:
            if str(p).startswith("/tmp/") and p.is_file():
                try:
                    os.remove(p)
                except Exception as e:
                    if verbose:
                        logger.error(f"Error removing temporary file {p}: {str(e)}")


def _download_if_url(path: Union[str, Path]) -> Union[str, Path]:
    """
    Downloads the file if the path is a URL and returns the local file path.
    """
    if isinstance(path, str) and urlparse(path).scheme in ("http", "https"):
        response = requests.get(path)
        response.raise_for_status()

        # Extract filename from URL or use a default one
        filename = Path(urlparse(path).path).name
        if not filename:
            filename = "downloaded_file"

        # Try to get extension from Content-Type header if not in filename
        if not Path(filename).suffix:
            content_type = response.headers.get("Content-Type")
            if content_type:
                extension = mimetypes.guess_extension(content_type.split(";")[0])
                if extension:
                    filename += extension

        # Add a unique identifier to the filename
        unique_filename = f"{uuid.uuid4()}_{filename}"
        local_path = Path("/tmp") / unique_filename
        with open(local_path, "wb") as file:
            file.write(response.content)
        return local_path
    return path


def _read_single_file(
    path: Union[str, Path],
    target: OutputType = "text",
    output: Union[Type[str], OutputFormat] = "document",
    verbose: bool = False,
) -> Union[Document, Dict, str, None]:
    """
    Reads a single file and returns its content based on the target format.
    """
    path = Path(path)
    mime_type, _ = mimetypes.guess_type(str(path))

    if mime_type is None:
        # Attempt to detect file type by reading the first few bytes
        with open(path, "rb") as f:
            header = f.read(5)
            if header == b"%PDF-":
                mime_type = "application/pdf"
            elif header[:2] == b"PK":
                # Possible DOCX or XLSX (which are zip files)
                # Read the content types in the file to distinguish
                f.seek(0)
                file_content = f.read()
                if b"word/" in file_content:
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif b"xl/" in file_content:
                    mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            elif (
                header.startswith(b"\xef\xbb\xbf")
                or header.startswith(b"\xfe\xff")
                or header.startswith(b"\xff\xfe")
            ):
                # Possible text file with BOM
                mime_type = "text/plain"
            else:
                # Check for JSON files
                try:
                    with open(path, "r", encoding="utf-8") as json_file:
                        json.load(json_file)
                    mime_type = "application/json"
                except (json.JSONDecodeError, UnicodeDecodeError):
                    # Default to binary if no match
                    mime_type = "application/octet-stream"

    try:
        content = None
        match mime_type:
            case "application/pdf":
                content = _read_pdf(path)
            case "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                content = _read_docx(path)
            case "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                content = _read_xlsx(path)
            case "text/csv":
                content = _read_csv(path)
            case "application/json":
                content = _read_json_content(
                    path
                )  # Use a function to read JSON content
            case _ if mime_type and mime_type.startswith("text/"):
                content = _read_text(path)
            case "application/xml":
                content = _read_xml(path)
            case _:
                content = _read_binary(path)

        if output == "json" and mime_type == "application/json":
            return content  # Return the JSON content directly

        metadata = {
            "file_name": path.name,
            "file_type": mime_type or "unknown",
            "file_size": path.stat().st_size,
        }

        if output == "document":
            return Document(content=content, metadata=metadata)
        elif output == str:
            return content

        return Document(content=content, metadata=metadata)

    except Exception as e:
        if verbose:
            logger.error(f"Error reading file {path}: {str(e)}")
        return None


def _read_json_content(path: Path) -> Dict:
    """
    Reads JSON files and returns their content as a dictionary.
    """
    try:
        with open(path, "r", encoding="utf-8") as file:
            return json.load(file)
    except Exception as e:
        logger.error(f"Error reading JSON {path}: {str(e)}")
        return {}


def _read_pdf(path: Path) -> str:
    """
    Extracts text from a PDF, including proper formatting for tables and paragraphs.
    """
    try:
        with open(path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                extracted_text = page.extract_text() or ""
                # Additional handling for simple table formatting
                # We'll assume columns are separated by enough space to infer table-like structures
                text += _format_pdf_text(extracted_text)
        return text.strip()
    except Exception as e:
        logger.error(f"Error reading PDF {path}: {str(e)}")
        return ""


def _format_pdf_text(extracted_text: str) -> str:
    """
    Post-process PDF extracted text to ensure proper formatting for tables, paragraphs, etc.
    This function attempts to identify and format tables, preserve paragraph structure,
    and handle common PDF extraction issues.
    """
    lines = extracted_text.split("\n")
    formatted_lines = []
    in_table = False
    table_column_widths = []

    for line in lines:
        stripped_line = line.strip()

        # Detect table start/end
        if "  " in line and not in_table:
            in_table = True
            table_column_widths = [len(col) for col in line.split()]
        elif in_table and not any(
            len(word) > width for word, width in zip(line.split(), table_column_widths)
        ):
            in_table = False

        if in_table:
            # Format table rows
            columns = line.split()
            formatted_line = " | ".join(
                col.ljust(width) for col, width in zip(columns, table_column_widths)
            )
            formatted_lines.append(formatted_line)
        elif stripped_line:
            # Handle regular text, attempting to preserve paragraph structure
            if formatted_lines and not formatted_lines[-1].endswith(
                (".", "!", "?", ":", ";")
            ):
                formatted_lines[-1] += " " + stripped_line
            else:
                formatted_lines.append(stripped_line)
        else:
            # Preserve empty lines between paragraphs
            if formatted_lines and formatted_lines[-1].strip():
                formatted_lines.append("")

    # Post-processing: remove redundant empty lines and fix common OCR issues
    cleaned_lines = []
    for line in formatted_lines:
        cleaned_line = line.replace(" |", "|").replace(
            "| ", "|"
        )  # Clean up table formatting
        cleaned_line = cleaned_line.replace("l", "I").replace(
            "0", "O"
        )  # Common OCR fixes
        if cleaned_line or (cleaned_lines and cleaned_lines[-1]):
            cleaned_lines.append(cleaned_line)

    return "\n".join(cleaned_lines)


def _read_docx(path: Path) -> str:
    """
    Reads text from a DOCX file, including tables.
    """
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Handle tables in the DOCX file
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                full_text.append(
                    " | ".join(row_data)
                )  # Simple table formatting with '|'

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading DOCX {path}: {str(e)}")
        return ""


def _read_xlsx(path: Path) -> List[List[str]]:
    """
    Reads an XLSX file and returns its data, including table formatting.
    """
    try:
        with openpyxl.load_workbook(path, read_only=True) as workbook:
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows():
                row_data = [
                    str(cell.value) if cell.value is not None else "" for cell in row
                ]
                data.append(row_data)
            return data
    except Exception as e:
        logger.error(f"Error reading XLSX {path}: {str(e)}")
        return []


def _read_csv(path: Path) -> List[List[str]]:
    """
    Reads CSV data.
    """
    try:
        with open(path, "r", newline="", encoding="utf-8") as csvfile:
            reader = csv.reader(csvfile)
            return list(reader)
    except Exception as e:
        logger.error(f"Error reading CSV {path}: {str(e)}")
        return []


def _read_text(path: Path) -> str:
    """
    Reads plain text files.
    """
    try:
        with open(path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading text file {path}: {str(e)}")
        return ""


def _read_xml(path: Path) -> Dict:
    """
    Reads XML files and converts them to a dictionary.
    """
    try:
        tree = ET.parse(path)
        root = tree.getroot()
        return _element_to_dict(root)
    except Exception as e:
        logger.error(f"Error reading XML {path}: {str(e)}")
        return {}


def _read_binary(path: Path) -> str:
    return f"Binary file: {path.name}"


def _element_to_dict(element):
    """
    Recursively converts XML elements to a dictionary.
    """
    result = {}
    for child in element:
        if len(child) == 0:
            result[child.tag] = child.text
        else:
            result[child.tag] = _element_to_dict(child)
    return result


def _to_markdown(content: Union[str, List, Dict]) -> str:
    """
    Converts the content into markdown format.
    """
    if isinstance(content, str):
        return content
    elif isinstance(content, list):
        return "\n".join([" | ".join(row) for row in content])
    elif isinstance(content, dict):
        return "\n".join([f"**{k}**: {v}" for k, v in content.items()])
    else:
        return str(content)
